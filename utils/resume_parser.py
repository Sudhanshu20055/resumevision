"""
resume_parser.py
Extracts raw text from an uploaded resume file (.pdf, .docx, .txt).
"""
import os
import pdfplumber
import docx


class ResumeParseError(Exception):
    """Raised when a resume file can't be read or parsed."""
    pass


def extract_text(filepath: str) -> str:
    """
    Extract plain text from a resume file based on its extension.

    Args:
        filepath: absolute path to the uploaded file.

    Returns:
        Extracted text as a single string.

    Raises:
        ResumeParseError: if the file type is unsupported or extraction fails,
                           or if no meaningful text could be extracted.
    """
    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext == ".pdf":
            text = _extract_from_pdf(filepath)
        elif ext == ".docx":
            text = _extract_from_docx(filepath)
        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            raise ResumeParseError(
                f"Unsupported file type '{ext}'. Please upload a PDF, DOCX, or TXT file."
            )
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError(f"Failed to read the file: {exc}") from exc

    text = text.strip()
    if len(text) < 30:
        raise ResumeParseError(
            "Couldn't extract meaningful text from this file. "
            "It may be a scanned/image-based document."
        )
    return text


def _extract_from_pdf(filepath: str) -> str:
    chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                chunks.append(page_text)
    return "\n".join(chunks)


def _extract_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables, since resumes sometimes use them for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)
